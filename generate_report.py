from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 样式辅助函数 ──────────────────────────────────────
def set_font(run, name='宋体', size=12, bold=False, color=None):
    run.font.name = name
    run._r.get_or_add_rPr()
    run._r.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    sizes = {1: 16, 2: 14, 3: 13}
    set_font(run, '黑体', sizes.get(level, 12), bold=True)
    return p

def add_para(doc, text, indent=0, size=12, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.first_line_indent = Pt(size * 2 * indent)
    run = p.add_run(text)
    set_font(run, '宋体', size, bold)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.space_before  = Pt(3)
    p.paragraph_format.space_after   = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run._r.get_or_add_rPr()
    run._r.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x6e)
    p.paragraph_format.first_line_indent = Pt(0)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, '黑体', 11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E1F2')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # 数据行
    for r_i, row in enumerate(rows):
        drow = table.rows[r_i+1]
        for c_i, val in enumerate(row):
            cell = drow.cells[c_i]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            set_font(run, '宋体', 10.5)
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

# ════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('数据结构课程设计')
set_font(r, '黑体', 26, bold=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('作业一：北邮本部校园导航系统')
set_font(r, '宋体', 16)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ('姓    名', '李韶庸'),
    ('学    院', ''),
    ('专    业', ''),
    ('班    级', ''),
    ('学    号', '2024213672'),
    ('班内序号', ''),
    ('指导教师', ''),
]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{label}：{val}{"＿＿＿＿" if not val else ""}')
    set_font(r, '宋体', 14)

doc.add_page_break()

# ════════════════════════════════════════════════════
# 一、实验题目
# ════════════════════════════════════════════════════
add_heading(doc, '一、实验题目', 1)
add_para(doc,
    '北邮本部校园导航系统',
    bold=True)
add_para(doc,
    '设计并实现一个基于图结构的校园导航系统，支持建筑检索、任意两点间最短路径查询、'
    '全校园建筑遍历规划，以及交互式地图可视化展示。系统数据来源于北邮本部真实校园地图，'
    '包含 172 个节点（建筑与路口）、133 条道路。', indent=1)

# ════════════════════════════════════════════════════
# 二、题目分析与算法设计
# ════════════════════════════════════════════════════
add_heading(doc, '二、题目分析与算法设计', 1)

add_heading(doc, '2.1 问题理解', 2)
add_para(doc,
    '本题本质是一个加权无向图上的多功能路径查询问题，具体包含三个子问题：', indent=1)

add_para(doc, '子问题一：最短路径查询', bold=True, indent=1)
add_para(doc,
    '给定起点 s 和终点 t，求图中从 s 到 t 的最短路径及其长度。这是经典的单源最短路径问题。', indent=1)

add_para(doc, '子问题二：全建筑遍历规划', bold=True, indent=1)
add_para(doc,
    '给定出发点，规划一条依次经过所有建筑（非路口节点）的访问顺序，使总行程尽可能短。'
    '本质上是旅行商问题（TSP）的近似版本，图中节点间并非直接相连，需结合最短路算法使用。', indent=1)

add_para(doc, '子问题三：建筑检索', bold=True, indent=1)
add_para(doc,
    '根据关键字在建筑名称和描述中进行多级模糊匹配，快速返回候选建筑列表。', indent=1)

add_para(doc, '数据规模：', bold=True, indent=1)
add_table(doc,
    ['项目', '数量'],
    [
        ['节点总数 V', '172（其中路口 36 个，建筑 136 个）'],
        ['边总数 E',   '133'],
        ['需遍历建筑数 N', '136'],
    ],
    col_widths=[5, 10]
)

# 2.2
add_heading(doc, '2.2 算法设计', 2)

add_para(doc, '（1）图的存储结构', bold=True, indent=1)
add_para(doc,
    '采用邻接表表示图，使用 Map<String, List<Edge>> 实现，以建筑名称字符串作为节点标识符。'
    '与邻接矩阵相比，邻接表在稀疏图（本题 E=133 远小于 V²=29584）下空间利用率更高，遍历邻居效率也更优。', indent=1)
add_code(doc, 'static class Edge { String target; int weight; }')
add_code(doc, 'Map<String, List<Edge>> adjList = new HashMap<>();')

add_para(doc, '（2）算法一：Dijkstra 最短路径', bold=True, indent=1)
add_para(doc,
    '使用优先队列优化的 Dijkstra 算法求单源最短路径：\n'
    '  ① 初始化：起点距离为 0，其余节点距离为正无穷\n'
    '  ② 每次从优先队列中取出当前距离最小的节点 u\n'
    '  ③ 松弛 u 的所有邻边：若 dist[u] + w(u,v) < dist[v]，更新 dist[v] 并记录前驱 prev[v]=u\n'
    '  ④ 重复直到队列为空，通过 prev 数组回溯重建完整路径\n'
    '时间复杂度：O((V+E)logV)', indent=1)
add_code(doc, 'PriorityQueue<String> pq = new PriorityQueue<>(')
add_code(doc, '    Comparator.comparingInt(n -> dist.getOrDefault(n, Integer.MAX_VALUE)));')

add_para(doc, '（3）算法二：最近邻贪心遍历', bold=True, indent=1)
add_para(doc,
    '全建筑遍历问题是 NP-Hard 的 TSP 问题，采用最近邻贪心策略给出近似解：\n'
    '  ① 从起点出发，将所有非路口建筑加入未访问集合\n'
    '  ② 每步对当前位置查询距离最近的未访问建筑\n'
    '  ③ 移动到该建筑，标记为已访问，重复直至全部访问完毕', indent=1)
add_para(doc, '优化：距离矩阵预计算缓存', bold=True, indent=1)
add_para(doc,
    '朴素实现中每步重新运行 Dijkstra，共需运行 N=136 次，总时间 O(N·(V+E)logV)。\n'
    '优化方案：首次调用遍历时，对所有 V=172 个节点预计算 Dijkstra 并缓存到距离矩阵。'
    '此后每次调用直接 O(1) 查表，总遍历开销降至 O(N²)。'
    '当图结构发生变更时，自动使缓存失效并在下次查询时重建。', indent=1)
add_code(doc, 'private synchronized void ensureDistMatrix() {')
add_code(doc, '    if (distMatrix != null) return;')
add_code(doc, '    for (String node : buildings.keySet())')
add_code(doc, '        distMatrix.put(node, dijkstra(node).dist);')
add_code(doc, '}')

add_para(doc, '（4）算法三：三级建筑检索', bold=True, indent=1)
add_para(doc, '对输入关键字按优先级分三级匹配，结果合并后返回前 12 条：', indent=1)
add_table(doc,
    ['优先级', '匹配规则', '示例'],
    [
        ['1（最高）', '名称完全等于关键字', '"图书馆" → 图书馆'],
        ['2',         '名称以关键字为前缀', '"图" → 图书馆南门'],
        ['3',         '名称或描述包含关键字', '"出入口" → 图书馆南门'],
    ],
    col_widths=[3, 7, 6]
)

add_para(doc, '（5）系统整体架构', bold=True, indent=1)
add_code(doc, '┌────────────────────────────────────────────────┐')
add_code(doc, '│  浏览器前端（index.html / annotate.html）       │')
add_code(doc, '│  Canvas 地图可视化 + 搜索框 + 结果展示          │')
add_code(doc, '└─────────────────┬──────────────────────────────┘')
add_code(doc, '                  │ HTTP / JSON API')
add_code(doc, '┌─────────────────▼──────────────────────────────┐')
add_code(doc, '│  Java HTTP Server（com.sun.net.httpserver）      │')
add_code(doc, '│  /api/route  /api/traverse  /api/buildings      │')
add_code(doc, '└─────────────────┬──────────────────────────────┘')
add_code(doc, '                  │')
add_code(doc, '┌─────────────────▼──────────────────────────────┐')
add_code(doc, '│  Dijkstra │ 最近邻贪心 │ 三级检索 │ 距离矩阵缓存│')
add_code(doc, '└─────────────────┬──────────────────────────────┘')
add_code(doc, '                  │')
add_code(doc, '┌─────────────────▼──────────────────────────────┐')
add_code(doc, '│  数据层（test.txt + coordinates.json）           │')
add_code(doc, '│  172 节点 / 133 条边 / 节点坐标                  │')
add_code(doc, '└────────────────────────────────────────────────┘')

# 2.3
add_heading(doc, '2.3 问题发现与解决', 2)

add_para(doc, '问题1：遍历功能响应慢', bold=True, indent=1)
add_para(doc,
    '初始实现每步重新运行 Dijkstra，136 次调用导致首次遍历有明显延迟。\n'
    '解决：引入距离矩阵缓存，首次调用时预计算所有节点的单源最短路径并存储，'
    '后续遍历直接查表，响应时间大幅缩短。', indent=1)

add_para(doc, '问题2：数据文件被非法输入破坏', bold=True, indent=1)
add_para(doc,
    '地图数据以逗号分隔的 CSV 格式存储，若建筑名称中包含逗号，写入后下次解析会产生错误分割，导致数据损坏。\n'
    '解决：在 addBuilding 中增加输入校验，名称或描述含逗号时拒绝操作并给出提示。', indent=1)

add_para(doc, '问题3：程序路径硬编码，换电脑无法运行', bold=True, indent=1)
add_para(doc,
    '原始代码以 "java" 字符串硬编码数据目录，依赖程序从项目根目录启动，换台电脑或换目录运行即失效。\n'
    '解决：实现三级路径自动检测（命令行参数 → 当前目录 → 从 class 文件位置向上查找），'
    '并提供跨平台启动脚本 run.sh（macOS/Linux）和 run.bat（Windows）。', indent=1)

# ════════════════════════════════════════════════════
# 三、复杂度分析
# ════════════════════════════════════════════════════
add_heading(doc, '三、程序复杂度分析', 1)

add_heading(doc, '3.1 时间复杂度', 2)
add_para(doc, '设图中节点数 V=172，边数 E=133，需遍历建筑数 N=136。', indent=1)
add_table(doc,
    ['功能', '时间复杂度', '说明'],
    [
        ['最短路径查询', 'O((V+E)logV)', '单次优先队列 Dijkstra'],
        ['遍历首次调用', 'O(V·(V+E)logV)', '预计算全量距离矩阵，V 次 Dijkstra'],
        ['遍历后续调用', 'O(N²)',          '直接查距离矩阵，N 步每步 O(N) 扫描'],
        ['建筑检索',     'O(V)',            '遍历所有建筑做字符串匹配'],
        ['增删建筑/道路', 'O(V+E)',         'CRUD + 缓存失效'],
    ],
    col_widths=[4, 4, 8]
)
add_para(doc,
    '对于本题规模，代入 V=172, E=133, N=136：单次最短路约 2257 次基本操作，几乎瞬时；'
    '遍历预计算约 38 万次操作，亦在毫秒级内完成。', indent=1)

add_heading(doc, '3.2 空间复杂度', 2)
add_table(doc,
    ['数据结构', '空间复杂度', '说明'],
    [
        ['邻接表 adjList',     'O(V+E)',  '存储所有节点及出边'],
        ['建筑列表 buildings', 'O(V)',    '节点名称与描述'],
        ['道路列表 roads',     'O(E)',    '用于持久化'],
        ['距离矩阵 distMatrix','O(V²)',   '每个节点到其余所有节点的距离'],
        ['Dijkstra 单次运行',  'O(V)',    '距离数组与优先队列'],
    ],
    col_widths=[5, 3.5, 7.5]
)
add_para(doc,
    '主要空间开销为距离矩阵，O(V²)=172²≈29584 个整数，约占 116 KB，内存开销极小。', indent=1)

# ════════════════════════════════════════════════════
# 四、运行结果展示
# ════════════════════════════════════════════════════
add_heading(doc, '四、程序运行结果展示', 1)
add_para(doc, '测试1：最短路径查询', bold=True, indent=1)
add_para(doc, '输入起点「学一食堂」，终点「图书馆南门」，系统显示最短路径列表、总距离，地图上绘制橙色路线。', indent=1)
add_para(doc, '【请在此处插入截图】', indent=1)
doc.add_paragraph()

add_para(doc, '测试2：全建筑遍历', bold=True, indent=1)
add_para(doc, '输入出发点「主楼东门」，系统输出建议访问 136 个建筑的顺序及总行程距离，地图绘制绿色路线。', indent=1)
add_para(doc, '【请在此处插入截图】', indent=1)
doc.add_paragraph()

add_para(doc, '测试3：建筑检索', bold=True, indent=1)
add_para(doc, '搜索关键字「宿舍」，下拉框展示前 12 个匹配建筑，按精确 > 前缀 > 包含排序。', indent=1)
add_para(doc, '【请在此处插入截图】', indent=1)
doc.add_paragraph()

add_para(doc, '测试4：非法输入校验', bold=True, indent=1)
add_para(doc, '在管理面板中新增名称为「测试,楼」的建筑，系统提示「名称不能包含逗号」，拒绝写入。', indent=1)
add_para(doc, '【请在此处插入截图】', indent=1)

# ════════════════════════════════════════════════════
# 五、代码提交说明
# ════════════════════════════════════════════════════
add_heading(doc, '五、代码提交说明', 1)
add_table(doc,
    ['项目', '内容'],
    [
        ['代码文件',
         'java/src/CampusNavSystem.java（后端主程序，含全部算法与 HTTP 服务）\n'
         'java/index.html（导航前端页面）\n'
         'java/annotate.html（坐标标注工具）'],
        ['编译/运行环境', 'JDK 17（兼容 JDK 8+），无第三方依赖，使用 JDK 内置 com.sun.net.httpserver'],
        ['编译命令', 'javac -encoding UTF-8 -d out/production/java java/src/CampusNavSystem.java'],
        ['运行方式', 'macOS/Linux：bash run.sh\nWindows：双击 run.bat\n脚本自动编译并启动，浏览器访问 http://localhost:8080'],
        ['输入数据说明', '地图数据：java/test.txt（172 节点、133 条边）\n节点坐标：java/coordinates.json（用于地图可视化）'],
    ],
    col_widths=[4, 12]
)

# ════════════════════════════════════════════════════
# 六、实验过程记录
# ════════════════════════════════════════════════════
add_heading(doc, '六、实验过程记录', 1)
add_table(doc,
    ['日期', '时长', '工作内容', '备注'],
    [
        ['＿月＿日', '2h', '分析题目要求，设计图的存储结构，选定邻接表方案', ''],
        ['＿月＿日', '3h', '实现 Dijkstra 算法与路径回溯，本地验证正确性', ''],
        ['＿月＿日', '2h', '实现最近邻贪心遍历，整合 HTTP 服务器框架', ''],
        ['＿月＿日', '3h', '完成前端页面，实现 Canvas 地图路径可视化', ''],
        ['＿月＿日', '1h', '标注 172 个节点的地图坐标，完善数据集', ''],
        ['＿月＿日', '2h', '发现并修复路径硬编码、逗号输入等问题，编写跨平台启动脚本', ''],
        ['＿月＿日', '1h', '优化遍历算法，引入距离矩阵缓存，测试全部功能', ''],
    ],
    col_widths=[3, 2, 9, 2]
)

# ════════════════════════════════════════════════════
# 七、心得与总结
# ════════════════════════════════════════════════════
add_heading(doc, '七、心得与总结', 1)
add_para(doc,
    '本次课程设计以北邮校园地图为背景，从零构建了一个完整的图论应用系统。在实践过程中有以下几点收获：', indent=1)

add_para(doc, '一、算法理解的深化。', bold=True, indent=1)
add_para(doc,
    'Dijkstra 算法在课堂上停留于理论层面，实际实现后才深刻体会到优先队列优化的意义——'
    '将朴素版本的 O(V²) 降低到 O((V+E)logV)，对中大规模图的性能影响显著。同时，'
    '通过预计算全量距离矩阵实现遍历提速，让我理解了「以空间换时间」这一经典权衡在工程中的具体体现。', indent=1)

add_para(doc, '二、数据结构选择的重要性。', bold=True, indent=1)
add_para(doc,
    '本题选用邻接表而非邻接矩阵，原因在于校园路网是稀疏图（E=133 << V²=29584），'
    '邻接表的空间与遍历效率均优于矩阵。这一选择在规模更大的地图数据上差距会更加明显。', indent=1)

add_para(doc, '三、系统设计的全局思维。', bold=True, indent=1)
add_para(doc,
    '一个完整的系统不仅仅是算法正确，还需要考虑数据完整性（逗号校验）、'
    '跨平台兼容性（路径自动检测、跨平台启动脚本）和用户体验（实时搜索、回车键支持、地图可视化）。'
    '这些工程细节虽不涉及核心算法，却决定了系统在真实场景中是否可用。', indent=1)

add_para(doc, '创新点：', bold=True, indent=1)
add_para(doc,
    '系统在遍历算法上引入了惰性距离矩阵缓存——首次调用时一次性预计算，后续调用直接复用，'
    '且在图结构变更时自动失效重建，兼顾了性能与正确性。前端采用 Canvas 实时绘制路径节点，'
    '辅以坐标标注工具，实现了完整的地图可视化闭环。', indent=1)

# ── 保存 ──────────────────────────────────────────
output_path = '实验报告_2024213672_李韶庸.docx'
doc.save(output_path)
print(f'✅ 报告已生成：{output_path}')
